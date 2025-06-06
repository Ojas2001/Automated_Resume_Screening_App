import streamlit as st
from docx import Document
from PyPDF2 import PdfReader
from langchain_mistralai import ChatMistralAI
from langchain_core.prompts import ChatPromptTemplate
import os
import docx
import base64
from datetime import datetime
from fpdf import FPDF
import tempfile
import re
import fitz
import unicodedata
from io import BytesIO
import io
import time


if "MISTRAL_API_KEY" not in os.environ:
    os.environ["MISTRAL_API_KEY"] = 'zQMwmcui0fmcy7QAYZfgKFeiXVZX6y2d'

llm = ChatMistralAI(
    model="mistral-large-latest",
    temperature=0,
    max_retries=2,
)

def get_file_type(file_path):
    try:
        Document(file_path)
        return "DOCX"
    except Exception:
        pass
    try:
        PdfReader(file_path)
        return "PDF"
    except Exception:
        pass
    return "Please upload PDF or WORD file only!!"

@st.cache_data
def clean_text(text: str) -> str:
    try:
        text = ''.join(ch for ch in text if unicodedata.category(ch) != 'Cf')
        text = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)
        text = text.replace("’", "'").replace("‘", "'") \
                .replace("“", '"').replace("”", '"') \
                .replace("–", "-").replace("—", "-")
        text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text)
        text = re.sub(r'[=_\-]{3,}', '', text)
        text = re.sub(r'\n+', '\n', text)  # collapse multiple newlines
        text = re.sub(r' +', ' ', text)    # collapse multiple spaces
        text = text.strip()                # remove leading/trailing spaces
        return text
    except Exception as e:
        st.error(f"Something wrong with the text!!: {e}")
        return None

# @st.cache_data
# def read_pdf(pdf_file):
#    try:
#        reader = PdfReader(pdf_file)
#         num_pages = len(reader.pages)
#        pdf_text_list = [] 
#         page = reader.pages
#        for i in range(num_pages):
#            pdf_text_list.append(page[i].extract_text())
#        pdf_text = '\n'.join(pdf_text_list)
#        return pdf_text
#    except Exception as e:
#        st.error(f"Error reading PDF file: {e}")
#        return None

@st.cache_data
def read_pdf(pdf_file):
    try:
        pdf_file.seek(0)
        pdf_bytes = pdf_file.read()
        pdf_stream = io.BytesIO(pdf_bytes)
        doc = fitz.open(stream=pdf_stream, filetype="pdf")
        text = ''
        for page in doc:
            text += page.get_text("text") + '\n'
        return text
    except Exception as e:
        st.error(f"Error reading PDF file: {e}")
        return None

@st.cache_data
def read_word(word_file):
    try:
        doc = docx.Document(word_file)
        word_text_list = []
        for para in doc.paragraphs:
            word_text_list.append(para.text)
        return '\n'.join(word_text_list)
    except Exception as e:
        st.error(f"Error reading Word file: {e}")
        return None

def run_llm(jd,resume):
    try:
        prompt = ChatPromptTemplate.from_messages(
            [
                (
                    "human",
                """Act as a recruiter evaluating whether a candidate’s resume aligns well with the given job description. Follow these steps:  
                1. Key Requirements: Identify the must-have skills, qualifications, and experience from the job description. Prioritize hard requirements like certifications, years of experience, technical skills, and so on.
                2. Resume Match: Analyze the resume to see which of these requirements are met. Note gaps. The candidate must meet **maximum and critical** requirements to qualify as a fit. Be stringent-ignore vague or loosely related skills.
                3. Feedback: Provide a concise but informative response with:  
                - Name of the candidate : "Candidate Name" \n
                - Job Role : "Indicate the Position Title or Job role mentioned in the job description" \n
                - Overall Fit: "Best fit" / "Good fit" / "Moderate fit" / "Poor fit" / "Bad Fit". \n
                - Matches: List key skills/experience the candidate has.  \n
                - Gaps: Mention critical missing qualifications (if any).  \n
                - Additional Notes: Highlight any transferable skills or notable achievements.  
                Please maintain a consistent format of output(Feedback only), given above.
                Keep the response crisp (3-5 lines max), avoid fluff, and focus on recruiter's needs. 
                Identify the critical requirements and required experience and then decide if the candiate is fit for the role or not.
                Keep the response as a clean text only with appropriate newlines, and no formatting is needed, like (** **) and so on.
                **Job Description**: {jd}  
                ---  
                **Resume**: {resume}  """
                )
            ]
        )

        prompt_1 = ChatPromptTemplate.from_messages(
            [
                (
                    "human",
                    """Act as a recruiter critically evaluating a candidate's resume against the job description. Follow these steps:
                    1. **Key Requirements**: Extract **non-negotiable skills, qualifications, and experience** from the job description. Prioritize hard requirements like certifications, years of experience, technical skills, and so on.
                    2. **Strict Resume Match**: The candidate must meet **maximum and critical** requirements to qualify as a fit. Be stringent-ignore vague or loosely related skills.
                    3. **Feedback**: Provide a structured, no-fluff response in the following format only:
                    - Candidate Name: [Candidate's Name]
                    - Job Role: [Extract Job role from  Job Description]
                    - Overall Fit: "Best Fit" (90-100% match) / "Strong Fit" (80-90%) / "Moderate Fit" (60-80%) / "Weak Fit" (40-60%) / "Reject" (<40%)
                    - Matches: List only **exact or close matches** (skill/experience directly from Job Description).
                    - Gaps: Highlight **critical deficiences** like missing certifications, insufficient years in a core skill, and so on.
                    - Additional Notes: Only mention **exceptional achievements** or **highly relevant transferable skills** that compensate for gaps.
                    Please maintain a consistent format of output(Feedback only), given above.
                    Keep the response crisp (3-5 lines max), avoid fluff, and focus on recruiter's needs. 
                    Keep the response as a clean text only with appropriate newlines, and no formatting is needed, like (** **) and so on.
                    **Job Description**: {jd}  
                    ---  
                    **Resume**: {resume}  """
                )
            ]
        )

        chain = prompt | llm
        run = chain.invoke(
            {
                "jd": jd,
                "resume": resume
            }
        )
        output = run.content
        candidate_name_line = [line for line in output.split("\n") if "Candidate Name:" in line][0]
        candidate_name = candidate_name_line.split("Candidate Name: ")[1].strip()
        job_role_line = [line for line in output.split("\n") if "Job Role:" in line][0]
        job_role = job_role_line.split("Job Role: ")[1].strip()
        fit_line = [line for line in output.split("\n") if "Overall Fit:" in line][0]
        fit_status = fit_line.split("Overall Fit: ")[1].strip()
        fit_colors = {
            "Best fit": "hsl(147, 50%, 47%)",
            "Good fit": "hsl(147, 50%, 50%)",
            "Moderate fit": "hsl(39, 100%, 50%)",
            "Poor fit": "hsl(0, 100%, 50%)",
            "Bad fit": "hsl(0, 100%, 25%)"
        }
        colored_fit_line = f'Overall Fit: <span style="color:{fit_colors[fit_status]}">{fit_status}</span>'
        colored_feedback = output.replace(fit_line, colored_fit_line)
        styled_feedback = f"""
            <div style="
                padding: 15px;
                border-radius: 10px;
                border: 1px solid #e0e0e0;
                margin: 10px 0;
                font-family: Arial, sans-serif;
            ">
            {colored_feedback}
            </div>
            """
        return styled_feedback, output, candidate_name, job_role
    except Exception:
        pass
    return "OOPS! Something went wrong!!"

def upload_files():
    file_types = ["pdf", "docx"]
    st.markdown("### 📄 **Upload Job Description file here (Only one at a time)**")
    jd_file = st.file_uploader("Upload Job Description file here (Only one at a time)", type=file_types, key="jd_file",label_visibility="collapsed")
    st.markdown("### 📄 **Upload Candidate Resume File here (One or Multiple)**")
    resume_files = st.file_uploader("Upload Candidate Resume File here (One or Multiple)", type=file_types, key="resume_file",label_visibility="collapsed",accept_multiple_files=True)
    if jd_file is not None and len(resume_files) > 0:
        st.success("Both files uploaded successfully!")
        return jd_file, resume_files
    return None, None

def show_pdf_download_button(pdf_path,candidate_name,job_role):
    """Display PDF download button"""
    with open(pdf_path, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
    
    download_link = (
        f'<a href="data:application/pdf;base64,{base64_pdf}" '
        f'download="resume_assessment_{candidate_name}_for_{job_role}.pdf">Download Report</a>'
    )
    st.markdown(download_link, unsafe_allow_html=True)
    st.success("PDF report generated successfully!")

def create_pdf_with_contents(candidate_name, job_role, output):
    """Create PDF with complete file contents"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=10)
    
    # Add title
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt=f"Assessment for {candidate_name} for {job_role}", ln=2, align='C')
    pdf.ln(10)
    
    # Process first file
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 10, txt="Result:", ln=1)
    pdf.set_font("Arial", size=10)
    
    pdf.multi_cell(0, 6, txt=output)
    pdf.ln(5)
    
    # Save to temporary file
    temp_dir = tempfile.gettempdir()
    pdf_path = os.path.join(temp_dir, f"resume_assessment_{candidate_name}_for_{job_role}.pdf")
    pdf.output(pdf_path)
    
    return pdf_path

def main():
    st.title("Automated Resume Screening Tool")
    jd_file,resume_files = upload_files()
    if jd_file is not None and len(resume_files) > 0:
            jd_file_type = get_file_type(jd_file)
            resume_file_types_list = []
            for r in resume_files:
                resume_file_types_list.append(get_file_type(r))

            if jd_file_type == "DOCX":
                jd_text = read_word(jd_file)
                jd_text = clean_text(jd_text)
            elif jd_file_type == "PDF":
                jd_text = read_pdf(jd_file)
                jd_text = clean_text(jd_text)
            else:
                st.error("OOPS! Job Description File can't be read.")
            
            resume_file_text_list = []
            for resume_file,r in zip(resume_files,resume_file_types_list):
                if r == "DOCX":
                    resume_text = read_word(resume_file)
                    resume_text = clean_text(resume_text)
                    resume_file_text_list.append(resume_text)
                elif r == "PDF":
                    resume_text = read_pdf(resume_file)
                    resume_text = clean_text(resume_text)
                    resume_file_text_list.append(resume_text)
                else:
                    st.error("OOPS! Resume File can't be read.")

            if st.button("Get Results"):
                for resume_file, resume_text in zip(resume_files, resume_file_text_list):
                    try:
                        result, output, candidate_name, job_role = run_llm(jd_text, resume_text)

                        with st.expander(f"📄 {candidate_name} - {resume_file.name}"):
                            st.markdown(result, unsafe_allow_html=True)
                            pdf_path = create_pdf_with_contents(candidate_name, job_role, output)
                            show_pdf_download_button(pdf_path, candidate_name, job_role)

                    except Exception:
                        with st.expander(f"📄 Error - {resume_file.name}"):
                            st.error("OOPS! There is some issue with the LLM API!!")
                    time.sleep(10)

    else:
        st.warning("Please upload both the files to proceed!!")

if __name__ == "__main__":
    main()